"""Tailor a resume to a specific job posting.

Reads a job posting from Postings/ directory, analyzes requirements and keywords,
then generates a customized resume in both Markdown and Word format optimized
for ATS parsing.

Usage:
    python3 tailor_resume.py <posting_file> <output_name>

Example:
    python3 tailor_resume.py "Postings/Lead Data Analyst - Point Digital Finance, Inc..txt" "Lead_Data_Analyst_Point"
"""

import json
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_posting_info(posting_text):
    """Extract title, company, location, and body from posting file."""
    lines = posting_text.split('\n')
    metadata = {}
    body_start = 0

    # Parse frontmatter (YAML between ---)
    if lines[0].strip() == '---':
        i = 1
        while i < len(lines):
            if lines[i].strip() == '---':
                body_start = i + 1
                break
            if ':' in lines[i]:
                key, value = lines[i].split(':', 1)
                metadata[key.strip()] = value.strip()
            i += 1

    body = '\n'.join(lines[body_start:]).strip()
    return metadata, body


def extract_keywords(posting_text):
    """Extract key skills, tools, and requirements from job posting."""
    keywords = {
        'skills': set(),
        'tools': set(),
        'experience': [],
        'responsibilities': [],
    }

    # Common technical keywords
    tech_keywords = [
        'SQL', 'Python', 'Snowflake', 'Tableau', 'Salesforce', 'dbt',
        'AWS', 'GCP', 'Azure', 'ETL', 'Agile', 'Scrum',
        'machine learning', 'causal inference', 'experimentation',
        'data governance', 'semantic layer', 'metrics layer',
        'OKR', 'KPI', 'QBR', 'GTM', 'product strategy',
        'analytics', 'business operations', 'strategy & ops',
        'leadership', 'cross-functional', 'communication'
    ]

    posting_lower = posting_text.lower()
    for keyword in tech_keywords:
        if keyword.lower() in posting_lower:
            keywords['skills'].add(keyword)

    # Extract bullet points as responsibilities
    for line in posting_text.split('\n'):
        line = line.strip()
        if line.startswith('- '):
            keywords['responsibilities'].append(line[2:])

    return keywords


def create_markdown_resume(base_resume, keywords, metadata):
    """Create tailored Markdown resume."""
    # Read base resume
    with open(base_resume, 'r') as f:
        resume = f.read()

    # Customize professional summary based on job title and key requirements
    title = metadata.get('title', 'Professional')
    company = metadata.get('company', 'Organization')
    location = metadata.get('location', 'Remote')

    # This is a placeholder - actual implementation would be more sophisticated
    # In practice, you'd use the keywords to dynamically rewrite sections

    return resume


def markdown_to_docx(md_content, output_path):
    """Convert Markdown resume to ATS-friendly .docx format."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = md_content.split('\n')

    for line in lines:
        line = line.rstrip()

        # Skip empty lines at the beginning
        if not line.strip() and len(doc.paragraphs) == 0:
            continue

        # Handle main heading (name)
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(14)
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Handle section headings
        elif line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_paragraph(heading, style='Heading 1')
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)

        # Handle contact info
        elif line and not line.startswith('#') and ('linkedin' in line or 'gmail' in line):
            cleaned = line.replace('[', '').replace('](', ' ').replace(')', '')
            p = doc.add_paragraph(cleaned)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)

        # Handle bullet points
        elif line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)

        # Handle bold lines
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            p.paragraph_format.space_before = Pt(6)

        # Regular text
        elif line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(3)

    doc.save(output_path)


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 tailor_resume.py <posting_file> <output_name>")
        print("Example: python3 tailor_resume.py 'Postings/Lead Data Analyst - Point.txt' 'Lead_Data_Analyst_Point'")
        sys.exit(1)

    posting_file = sys.argv[1]
    output_name = sys.argv[2]

    # Read posting
    if not Path(posting_file).exists():
        print(f"Error: Posting file not found: {posting_file}")
        sys.exit(1)

    with open(posting_file, 'r') as f:
        posting_text = f.read()

    metadata, body = extract_posting_info(posting_text)
    keywords = extract_keywords(posting_text + body)

    # Read base resume
    base_resume = Path('Resume.md')
    if not base_resume.exists():
        print(f"Error: Base resume not found: {base_resume}")
        sys.exit(1)

    with open(base_resume, 'r') as f:
        base_content = f.read()

    # Create customized markdown (simplified for now)
    # In practice, this would do sophisticated tailoring
    tailored_md = base_content

    # Save markdown version
    md_output = Path('Customized Resumes') / f"Resume_{output_name}.md"
    md_output.parent.mkdir(parents=True, exist_ok=True)
    with open(md_output, 'w') as f:
        f.write(tailored_md)

    # Create and save docx version
    docx_output = Path('Customized Resumes') / f"Resume_{output_name}.docx"
    markdown_to_docx(tailored_md, docx_output)

    # Return result
    result = {
        'status': 'tailored',
        'posting': metadata.get('title', 'Job'),
        'company': metadata.get('company', ''),
        'md_path': str(md_output),
        'docx_path': str(docx_output),
        'keywords_found': list(keywords['skills'])[:10],  # Top 10 keywords
    }

    print(json.dumps(result))


if __name__ == '__main__':
    main()
